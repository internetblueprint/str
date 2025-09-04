import streamlit as st
from datetime import datetime, timedelta
import re
from io import BytesIO
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# For PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus.tableofcontents import TableOfContents
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# For Word generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def add_signature_section_to_docx(doc, contract_data):
    """Add signature section to Word document"""
    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement on the date first written above.")
    doc.add_paragraph()
    
    doc.add_paragraph("SIGNED AT _________________________ ON _________________________")
    doc.add_paragraph()
    
    # Company signature
    doc.add_paragraph("THE COMPANY:")
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph(f"{contract_data['company_rep']}")
    doc.add_paragraph(f"{contract_data['company_position']}")
    doc.add_paragraph(f"{contract_data['company_name']}")
    doc.add_paragraph()
    
    doc.add_paragraph("WITNESS:")
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph("Full Name:")
    doc.add_paragraph("Date:")
    doc.add_paragraph()
    
    # Recipient signature
    doc.add_paragraph("THE RECIPIENT:")
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph(f"{contract_data['other_name']}")
    if contract_data['other_id']:
        doc.add_paragraph(f"ID Number: {contract_data['other_id']}")
    doc.add_paragraph()
    
    doc.add_paragraph("WITNESS:")
    doc.add_paragraph()
    doc.add_paragraph("_________________________________")
    doc.add_paragraph("Full Name:")
    doc.add_paragraph("Date:")
    doc.add_paragraph()
    
    # Legal disclaimer
    doc.add_page_break()
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run("LEGAL DISCLAIMER:")
    disclaimer_run.bold = True
    disclaimer_para.add_run(" This NDA has been generated to comply with South African law as of 2024. However, legal requirements may change, and specific circumstances may require additional provisions. It is recommended to have this agreement reviewed by a qualified South African attorney before execution.")

def amount_in_words(amount):
    """Convert numeric amount to words (enhanced version)"""
    if amount == 0:
        return "Zero"
    
    # Handle thousands
    if amount < 1000:
        return str(amount)
    elif amount < 10000:
        thousands = amount // 1000
        remainder = amount % 1000
        if remainder == 0:
            return f"{thousands} Thousand"
        else:
            return f"{thousands} Thousand {remainder}"
    elif amount < 100000:
        return f"{amount // 1000} Thousand"
    elif amount < 1000000:
        thousands = amount // 1000
        if thousands % 100 == 0:
            return f"{thousands // 100} Hundred Thousand"
        else:
            return f"{thousands} Thousand"
    else:
        # Handle millions
        millions = amount // 1000000
        remainder = amount % 1000000
        if remainder == 0:
            return f"{millions} Million"
        elif remainder < 1000:
            return f"{millions} Million {remainder}"
        else:
            thousands = remainder // 1000
            if thousands == 0:
                return f"{millions} Million"
            else:
                return f"{millions} Million {thousands} Thousand"

def main():
    st.set_page_config(
        page_title="SA NDA Generator",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("🇿🇦 South African Non-Disclosure Agreement Generator")
    st.markdown("*Compliant with SA Constitution, LRA, BCEA, POPIA, Competition Act & Protected Disclosures Act*")
    
    # Check for dependencies
    if not PDF_AVAILABLE:
        st.warning("⚠️ PDF export not available. Install reportlab: `pip install reportlab`")
    if not DOCX_AVAILABLE:
        st.warning("⚠️ Word export not available. Install python-docx: `pip install python-docx`")
    
    # Sidebar for contract type
    st.sidebar.header("Contract Configuration")
    contract_type = st.sidebar.selectbox(
        "Select NDA Type:",
        ["Employee NDA", "Contractor NDA", "Mutual NDA"]
    )
    
    # Export format selection
    st.sidebar.header("Export Options")
    export_formats = ["Text (.txt)"]
    if PDF_AVAILABLE:
        export_formats.append("PDF (.pdf)")
    if DOCX_AVAILABLE:
        export_formats.append("Word (.docx)")
    
    export_format = st.sidebar.selectbox("Download Format:", export_formats)
    
    # Main form
    col1, col2 = st.columns(2)
    
    with col1:
        st.header("Party Details")
        
        # Company details
        st.subheader("🏢 Disclosing Party (Company)")
        company_name = st.text_input("Company Name*", placeholder="ABC (Pty) Ltd")
        company_reg = st.text_input("Registration Number*", placeholder="2023/123456/07")
        company_address = st.text_area("Registered Address*", placeholder="123 Main Street, Johannesburg, 2001")
        company_rep = st.text_input("Authorized Representative*", placeholder="John Smith")
        company_position = st.text_input("Representative Position*", placeholder="Managing Director")
        
        # Individual/Other party details
        st.subheader("👤 Receiving Party")
        if contract_type == "Mutual NDA":
            other_party_type = st.selectbox("Other Party Type:", ["Company", "Individual"])
            if other_party_type == "Company":
                other_name = st.text_input("Company Name*", placeholder="XYZ (Pty) Ltd")
                other_reg = st.text_input("Registration Number*", placeholder="2023/654321/07")
                other_address = st.text_area("Registered Address*", placeholder="456 Business Ave, Cape Town, 8001")
            else:
                other_name = st.text_input("Full Name*", placeholder="Jane Doe")
                other_id = st.text_input("ID Number*", placeholder="8501015800083")
                other_address = st.text_area("Residential Address*", placeholder="789 Residential St, Durban, 4001")
        else:
            other_name = st.text_input("Full Name*", placeholder="Jane Doe")
            other_id = st.text_input("ID Number*", placeholder="8501015800083")
            other_address = st.text_area("Residential Address*", placeholder="789 Residential St, Durban, 4001")
            if contract_type == "Employee NDA":
                job_title = st.text_input("Job Title/Position*", placeholder="Software Developer")
                employment_date = st.date_input("Employment Start Date*", datetime.now())
    
    with col2:
        st.header("NDA Terms")
        
        # Confidential information definition
        st.subheader("🔒 Confidential Information")
        confidential_info = st.multiselect(
            "Select types of confidential information:",
            [
                "Technical information and trade secrets",
                "Business strategies and plans",
                "Customer lists and client information",
                "Financial information and pricing",
                "Software source code and algorithms",
                "Marketing strategies and campaigns",
                "Supplier and vendor information",
                "Research and development data",
                "Personnel information (subject to POPIA)",
                "Manufacturing processes and methods"
            ],
            default=["Technical information and trade secrets", "Business strategies and plans"]
        )
        
        additional_info = st.text_area(
            "Additional confidential information (optional):",
            placeholder="Specify any additional confidential information..."
        )
        
        # Duration and scope
        st.subheader("⏰ Duration & Scope")
        duration_years = st.selectbox("Confidentiality Duration (years):", [1, 2, 3, 5, 10, "Indefinite"])
        
        if contract_type != "Mutual NDA":
            geographic_scope = st.selectbox(
                "Geographic Scope:",
                ["South Africa only", "Africa", "Global"]
            )
            
            post_employment = st.checkbox(
                "Extends beyond employment termination",
                value=True
            )
        
        # Remedies
        st.subheader("⚖️ Remedies for Breach")
        liquidated_damages = st.checkbox("Include liquidated damages clause")
        if liquidated_damages:
            damages_amount = st.number_input("Liquidated Damages Amount (ZAR):", min_value=0, value=50000, step=5000)
        
        interdict_relief = st.checkbox("Include interdict/injunctive relief", value=True)
        
        # POPIA compliance
        st.subheader("🛡️ POPIA Compliance")
        involves_personal_data = st.checkbox("Involves processing of personal information")
        if involves_personal_data:
            data_types = st.multiselect(
                "Types of personal information:",
                ["Employee personal data", "Customer personal data", "Supplier personal data", "Other personal data"]
            )
    
    # Generate contract button
    if st.button("Generate NDA Contract", type="primary"):
        if validate_inputs(company_name, company_reg, company_address, other_name):
            contract_data = {
                'contract_type': contract_type,
                'company_name': company_name,
                'company_reg': company_reg,
                'company_address': company_address,
                'company_rep': company_rep,
                'company_position': company_position,
                'other_name': other_name,
                'other_id': other_id if contract_type != "Mutual NDA" or (contract_type == "Mutual NDA" and other_party_type == "Individual") else None,
                'other_address': other_address,
                'confidential_info': confidential_info,
                'additional_info': additional_info,
                'duration_years': duration_years,
                'geographic_scope': geographic_scope if contract_type != "Mutual NDA" else "South Africa only",
                'liquidated_damages': liquidated_damages,
                'damages_amount': damages_amount if liquidated_damages else 0,
                'interdict_relief': interdict_relief,
                'involves_personal_data': involves_personal_data,
                'data_types': data_types if involves_personal_data else [],
                'job_title': job_title if contract_type == "Employee NDA" else None,
                'employment_date': employment_date if contract_type == "Employee NDA" else None,
                'post_employment': post_employment if contract_type != "Mutual NDA" else True
            }
            
            st.success("✅ NDA Contract Generated Successfully!")
            
            # Generate and display based on format
            if export_format == "Text (.txt)":
                contract_text = generate_nda_text(contract_data)
                st.subheader("📄 Generated NDA Contract")
                st.text_area("Contract Text:", value=contract_text, height=400)
                
                st.download_button(
                    label="📥 Download as Text",
                    data=contract_text,
                    file_name=f"SA_NDA_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain"
                )
            
            elif export_format == "PDF (.pdf)" and PDF_AVAILABLE:
                pdf_buffer = generate_nda_pdf(contract_data)
                st.subheader("📄 Generated NDA Contract (PDF Preview)")
                st.success("PDF generated successfully! Use the download button below.")
                
                st.download_button(
                    label="📥 Download as PDF",
                    data=pdf_buffer.getvalue(),
                    file_name=f"SA_NDA_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf"
                )
            
            elif export_format == "Word (.docx)" and DOCX_AVAILABLE:
                docx_buffer = generate_nda_docx(contract_data)
                st.subheader("📄 Generated NDA Contract (Word Preview)")
                st.success("Word document generated successfully! Use the download button below.")
                
                st.download_button(
                    label="📥 Download as Word Document",
                    data=docx_buffer.getvalue(),
                    file_name=f"SA_NDA_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("❌ Please fill in all required fields marked with *")

def validate_inputs(company_name, company_reg, company_address, other_name):
    return all([company_name, company_reg, company_address, other_name])

def generate_nda_pdf(contract_data):
    """Generate PDF version of the NDA with proper formatting"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Create custom styles
    styles = getSampleStyleSheet()
    
    # Custom title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1,  # Center alignment
        textColor=colors.black,
        fontName='Helvetica-Bold'
    )
    
    # Custom heading styles
    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.black,
        fontName='Helvetica-Bold'
    )
    
    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=15,
        textColor=colors.black,
        fontName='Helvetica-Bold'
    )
    
    # Body text style
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        alignment=4,  # Justify
        fontName='Helvetica'
    )
    
    # Build the document
    story = []
    current_date = datetime.now().strftime("%d %B %Y")
    
    # Title
    story.append(Paragraph("NON-DISCLOSURE AGREEMENT", title_style))
    story.append(Paragraph("(Compliant with South African Law)", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Agreement details
    story.append(Paragraph(f"THIS AGREEMENT is made on {current_date}", body_style))
    story.append(Spacer(1, 12))
    
    # Parties
    story.append(Paragraph("BETWEEN:", h2_style))
    
    # Company details
    company_text = f"""(1) <b>{contract_data['company_name'].upper()}</b> (Registration Number: {contract_data['company_reg']}), 
    a company duly incorporated in accordance with the laws of the Republic of South Africa, with its registered address at 
    {contract_data['company_address']} (hereinafter referred to as "the Company" or "Disclosing Party"), represented herein by 
    {contract_data['company_rep']}, {contract_data['company_position']}, who warrants that he/she has the necessary authority to bind the Company; and"""
    
    story.append(Paragraph(company_text, body_style))
    story.append(Spacer(1, 8))
    
    # Other party details
    other_party_text = f"""(2) <b>{contract_data['other_name'].upper()}</b>{f", ID Number: {contract_data['other_id']}" if contract_data['other_id'] else ""}, 
    with address at {contract_data['other_address']} (hereinafter referred to as "the Recipient" or "Receiving Party")
    {f", employed as {contract_data['job_title']} from {contract_data['employment_date'].strftime('%d %B %Y')}" if contract_data['contract_type'] == 'Employee NDA' else ""}."""
    
    story.append(Paragraph(other_party_text, body_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("(The Company and the Recipient may be referred to individually as a \"Party\" and collectively as the \"Parties\")", body_style))
    story.append(Spacer(1, 20))
    
    # RECITALS
    story.append(Paragraph("RECITALS", h1_style))
    
    recitals = [
        "WHEREAS, the Company possesses certain confidential and proprietary information, trade secrets, and intellectual property that constitute valuable business assets;",
        f"WHEREAS, the Recipient {'is employed by' if contract_data['contract_type'] == 'Employee NDA' else 'will be engaged by'} the Company and will have access to such confidential information in the course of {'employment' if contract_data['contract_type'] == 'Employee NDA' else 'the engagement'};",
        "WHEREAS, the Parties wish to protect the confidentiality of such information in accordance with the laws of the Republic of South Africa, including but not limited to the Constitution of South Africa (1996), Labour Relations Act 66 of 1995, Basic Conditions of Employment Act 75 of 1997, Protection of Personal Information Act 4 of 2013, Competition Act 89 of 1998, and Protected Disclosures Act 26 of 2000;"
    ]
    
    for recital in recitals:
        story.append(Paragraph(recital, body_style))
        story.append(Spacer(1, 8))
    
    story.append(Paragraph("NOW THEREFORE, the Parties agree as follows:", body_style))
    story.append(Spacer(1, 20))
    
    # Main clauses
    add_main_clauses_to_story(story, contract_data, h1_style, h2_style, body_style)
    
    # Signature section
    add_signature_section_to_story(story, contract_data, h1_style, body_style)
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_nda_docx(contract_data):
    """Generate Word document version of the NDA"""
    buffer = BytesIO()
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    heading1_style = doc.styles['Heading 1']
    heading2_style = doc.styles['Heading 2']
    normal_style = doc.styles['Normal']
    
    # Title
    title = doc.add_paragraph("NON-DISCLOSURE AGREEMENT", style=title_style)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("(Compliant with South African Law)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Space
    
    current_date = datetime.now().strftime("%d %B %Y")
    doc.add_paragraph(f"THIS AGREEMENT is made on {current_date}")
    doc.add_paragraph()
    
    # Parties section
    doc.add_heading("BETWEEN:", level=2)
    
    # Company details
    company_para = doc.add_paragraph()
    company_para.add_run(f"(1) {contract_data['company_name'].upper()}").bold = True
    company_para.add_run(f" (Registration Number: {contract_data['company_reg']}), a company duly incorporated in accordance with the laws of the Republic of South Africa, with its registered address at {contract_data['company_address']} (hereinafter referred to as \"the Company\" or \"Disclosing Party\"), represented herein by {contract_data['company_rep']}, {contract_data['company_position']}, who warrants that he/she has the necessary authority to bind the Company; and")
    
    # Other party details
    other_para = doc.add_paragraph()
    other_para.add_run(f"(2) {contract_data['other_name'].upper()}").bold = True
    other_para.add_run(f"{f', ID Number: {contract_data['other_id']}' if contract_data['other_id'] else ''}, with address at {contract_data['other_address']} (hereinafter referred to as \"the Recipient\" or \"Receiving Party\"){f', employed as {contract_data['job_title']} from {contract_data['employment_date'].strftime('%d %B %Y')}' if contract_data['contract_type'] == 'Employee NDA' else ''}.")
    
    doc.add_paragraph("(The Company and the Recipient may be referred to individually as a \"Party\" and collectively as the \"Parties\")")
    doc.add_paragraph()
    
    # RECITALS
    doc.add_heading("RECITALS", level=1)
    
    recitals = [
        "WHEREAS, the Company possesses certain confidential and proprietary information, trade secrets, and intellectual property that constitute valuable business assets;",
        f"WHEREAS, the Recipient {'is employed by' if contract_data['contract_type'] == 'Employee NDA' else 'will be engaged by'} the Company and will have access to such confidential information in the course of {'employment' if contract_data['contract_type'] == 'Employee NDA' else 'the engagement'};",
        "WHEREAS, the Parties wish to protect the confidentiality of such information in accordance with the laws of the Republic of South Africa, including but not limited to the Constitution of South Africa (1996), Labour Relations Act 66 of 1995, Basic Conditions of Employment Act 75 of 1997, Protection of Personal Information Act 4 of 2013, Competition Act 89 of 1998, and Protected Disclosures Act 26 of 2000;"
    ]
    
    for recital in recitals:
        doc.add_paragraph(recital)
    
    doc.add_paragraph("NOW THEREFORE, the Parties agree as follows:")
    doc.add_paragraph()
    
    # Add main content
    add_main_clauses_to_docx(doc, contract_data)
    
    # Add signature section
    add_signature_section_to_docx(doc, contract_data)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def add_main_clauses_to_story(story, contract_data, h1_style, h2_style, body_style):
    """Add main contract clauses to PDF story"""
    # Section 1: Definition of Confidential Information
    story.append(Paragraph("1. DEFINITION OF CONFIDENTIAL INFORMATION", h1_style))
    
    story.append(Paragraph("1.1 \"Confidential Information\" shall mean all non-public, proprietary, or confidential information disclosed by the Company to the Recipient, whether orally, in writing, electronically, or by observation, including but not limited to:", body_style))
    
    for i, item in enumerate(contract_data['confidential_info'], 1):
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{i}. {item};", body_style))
    
    if contract_data['additional_info']:
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{len(contract_data['confidential_info']) + 1}. {contract_data['additional_info']};", body_style))
    
    # Exceptions
    story.append(Paragraph("1.2 Confidential Information shall not include information that:", body_style))
    exceptions = [
        "Is or becomes publicly available through no breach of this Agreement by the Recipient;",
        "Was rightfully known by the Recipient before disclosure by the Company;",
        "Is rightfully received by the Recipient from a third party without breach of any confidentiality obligation;",
        "Is independently developed by the Recipient without use of or reference to the Confidential Information;",
        "Is required to be disclosed by law, regulation, or court order, provided that the Recipient gives the Company reasonable advance notice of such requirement."
    ]
    
    for i, exception in enumerate(exceptions, ord('a')):
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;({chr(i)}) {exception}", body_style))
    
    # Continue with other sections
    add_remaining_clauses_to_story(story, contract_data, h1_style, h2_style, body_style)

def add_main_clauses_to_docx(doc, contract_data):
    """Add main contract clauses to Word document"""
    # Section 1: Definition of Confidential Information
    doc.add_heading("1. DEFINITION OF CONFIDENTIAL INFORMATION", level=1)
    
    doc.add_paragraph("1.1 \"Confidential Information\" shall mean all non-public, proprietary, or confidential information disclosed by the Company to the Recipient, whether orally, in writing, electronically, or by observation, including but not limited to:")
    
    for i, item in enumerate(contract_data['confidential_info'], 1):
        doc.add_paragraph(f"{i}. {item};", style='List Number')
    
    if contract_data['additional_info']:
        doc.add_paragraph(f"{len(contract_data['confidential_info']) + 1}. {contract_data['additional_info']};", style='List Number')
    
    # Exceptions
    doc.add_paragraph("1.2 Confidential Information shall not include information that:")
    exceptions = [
        "Is or becomes publicly available through no breach of this Agreement by the Recipient;",
        "Was rightfully known by the Recipient before disclosure by the Company;",
        "Is rightfully received by the Recipient from a third party without breach of any confidentiality obligation;",
        "Is independently developed by the Recipient without use of or reference to the Confidential Information;",
        "Is required to be disclosed by law, regulation, or court order, provided that the Recipient gives the Company reasonable advance notice of such requirement."
    ]
    
    for i, exception in enumerate(exceptions, ord('a')):
        doc.add_paragraph(f"({chr(i)}) {exception}", style='List Bullet')
    
    # Continue with other sections
    add_remaining_clauses_to_docx(doc, contract_data)

def add_remaining_clauses_to_story(story, contract_data, h1_style, h2_style, body_style):
    """Add remaining clauses to PDF"""
    # Note: This is a placeholder in the original code. Add actual clauses here.
    story.append(Paragraph("2. OBLIGATIONS OF THE RECIPIENT", h1_style))
    story.append(Paragraph("3. CONSTITUTIONAL AND STATUTORY COMPLIANCE", h1_style))
    # TODO: Implement additional clauses (e.g., obligations, compliance, remedies, etc.)
    # Example placeholder clauses:
    story.append(Paragraph("2.1 The Recipient shall not disclose, publish, or otherwise reveal any of the Confidential Information received from the Company to any other party whatsoever except with the specific prior written authorization of the Company.", body_style))
    story.append(Paragraph("3.1 This Agreement complies with the Protection of Personal Information Act 4 of 2013 (POPIA) and other relevant South African legislation.", body_style))

def add_remaining_clauses_to_docx(doc, contract_data):
    """Add remaining clauses to Word document"""
    # Note: This is a placeholder in the original code. Add actual clauses here.
    doc.add_heading("2. OBLIGATIONS OF THE RECIPIENT", level=1)
    doc.add_heading("3. CONSTITUTIONAL AND STATUTORY COMPLIANCE", level=1)
    # TODO: Implement additional clauses (e.g., obligations, compliance, remedies, etc.)
    # Example placeholder clauses:
    doc.add_paragraph("2.1 The Recipient shall not disclose, publish, or otherwise reveal any of the Confidential Information received from the Company to any other party whatsoever except with the specific prior written authorization of the Company.")
    doc.add_paragraph("3.1 This Agreement complies with the Protection of Personal Information Act 4 of 2013 (POPIA) and other relevant South African legislation.")

def add_signature_section_to_story(story, contract_data, h1_style, body_style):
    """Add signature section to PDF"""
    story.append(Spacer(1, 20))
    story.append(Paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement on the date first written above.", body_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("SIGNED AT _________________________ ON _________________________", body_style))
    story.append(Spacer(1, 20))
    
    # Company signature
    story.append(Paragraph("THE COMPANY:", body_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("_________________________________", body_style))
    story.append(Paragraph(f"{contract_data['company_rep']}", body_style))
    story.append(Paragraph(f"{contract_data['company_position']}", body_style))
    story.append(Paragraph(f"{contract_data['company_name']}", body_style))
    story.append(Spacer(1, 20))
    
    # Witness
    story.append(Paragraph("WITNESS:", body_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("_________________________________", body_style))
    story.append(Paragraph("Full Name:", body_style))
    story.append(Paragraph("Date:", body_style))
    story.append(Spacer(1, 20))
    
    # Recipient signature
    story.append(Paragraph("THE RECIPIENT:", body_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("_________________________________", body_style))
    story.append(Paragraph(f"{contract_data['other_name']}", body_style))
    if contract_data['other_id']:
        story.append(Paragraph(f"ID Number: {contract_data['other_id']}", body_style))
    story.append(Spacer(1, 20))
    
    # Witness
    story.append(Paragraph("WITNESS:", body_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("_________________________________", body_style))
    story.append(Paragraph("Full Name:", body_style))
    story.append(Paragraph("Date:", body_style))
    story.append(Spacer(1, 20))
    
    # Legal disclaimer
    story.append(Paragraph("LEGAL DISCLAIMER:", h1_style))
    story.append(Paragraph("This NDA has been generated to comply with South African law as of 2024. However, legal requirements may change, and specific circumstances may require additional provisions. It is recommended to have this agreement reviewed by a qualified South African attorney before execution.", body_style))

def generate_nda_text(contract_data):
    """Generate plain text version of the NDA"""
    current_date = datetime.now().strftime("%d %B %Y")
    
    # Determine duration text
    if contract_data['duration_years'] == "Indefinite":
        duration_text = "indefinitely"
        duration_clause = "This obligation shall survive indefinitely"
    else:
        duration_text = f"{contract_data['duration_years']} years"
        duration_clause = f"This obligation shall survive for a period of {contract_data['duration_years']} years"
    
    if contract_data['post_employment'] and contract_data['contract_type'] == "Employee NDA":
        duration_clause += " from the termination of employment"
    elif contract_data['contract_type'] == "Contractor NDA":
        duration_clause += " from the completion or termination of the contractual relationship"
    else:
        duration_clause += " from the date of this Agreement"
    
    # Build confidential information list
    conf_info_list = []
    for i, item in enumerate(contract_data['confidential_info'], 1):
        conf_info_list.append(f"        {i}. {item};")
    
    if contract_data['additional_info']:
        conf_info_list.append(f"        {len(contract_data['confidential_info']) + 1}. {contract_data['additional_info']};")
    
    conf_info_text = "\n".join(conf_info_list)
    
    # Corrected recipient text
    recipient_text = f"(2) {contract_data['other_name'].upper()}{f', ID Number: {contract_data['other_id']}' if contract_data['other_id'] else ''}, with address at {contract_data['other_address']} (hereinafter referred to as \"the Recipient\" or \"Receiving Party\"){f', employed as {contract_data['job_title']} from {contract_data['employment_date'].strftime('%d %B %Y')}' if contract_data['contract_type'] == 'Employee NDA' else ''}."

    contract = f"""
NON-DISCLOSURE AGREEMENT
(Compliant with South African Law)

THIS AGREEMENT is made on {current_date}

BETWEEN:

(1) {contract_data['company_name'].upper()} (Registration Number: {contract_data['company_reg']}), a company duly incorporated in accordance with the laws of the Republic of South Africa, with its registered address at {contract_data['company_address']} (hereinafter referred to as "the Company" or "Disclosing Party"), represented herein by {contract_data['company_rep']}, {contract_data['company_position']}, who warrants that he/she has the necessary authority to bind the Company; and

{recipient_text}

(The Company and the Recipient may be referred to individually as a "Party" and collectively as the "Parties")

RECITALS

WHEREAS, the Company possesses certain confidential and proprietary information, trade secrets, and intellectual property that constitute valuable business assets;

WHEREAS, the Recipient {f"is employed by" if contract_data['contract_type'] == "Employee NDA" else "will be engaged by"} the Company and will have access to such confidential information in the course of {"employment" if contract_data['contract_type'] == "Employee NDA" else "the engagement"};

WHEREAS, the Parties wish to protect the confidentiality of such information in accordance with the laws of the Republic of South Africa, including but not limited to the Constitution of South Africa (1996), Labour Relations Act 66 of 1995, Basic Conditions of Employment Act 75 of 1997, Protection of Personal Information Act 4 of 2013, Competition Act 89 of 1998, and Protected Disclosures Act 26 of 2000;

NOW THEREFORE, the Parties agree as follows:

1. DEFINITION OF CONFIDENTIAL INFORMATION

1.1 "Confidential Information" shall mean all non-public, proprietary, or confidential information disclosed by the Company to the Recipient, whether orally, in writing, electronically, or by observation, including but not limited to:

{conf_info_text}

1.2 Confidential Information shall not include information that:
    (a) Is or becomes publicly available through no breach of this Agreement by the Recipient;
    (b) Was rightfully known by the Recipient before disclosure by the Company;
    (c) Is rightfully received by the Recipient from a third party without breach of any confidentiality obligation;
    (d) Is independently developed by the Recipient without use of or reference to the Confidential Information;
    (e) Is required to be disclosed by law, regulation, or court order, provided that the Recipient gives the Company reasonable advance notice of such requirement.

2. OBLIGATIONS OF THE RECIPIENT

2.1 The Recipient acknowledges that the Confidential Information is proprietary to the Company and constitutes valuable trade secrets.

IN WITNESS WHEREOF, the Parties have executed this Agreement on the date first written above.

SIGNED AT _________________________ ON _________________________

THE COMPANY:
_________________________________
{contract_data['company_rep']}
{contract_data['company_position']}
{contract_data['company_name']}

WITNESS:
_________________________________
Full Name:
Date:

THE RECIPIENT:
_________________________________
{contract_data['other_name']}
{f"ID Number: {contract_data['other_id']}" if contract_data['other_id'] else ''}

WITNESS:
_________________________________
Full Name:
Date:

LEGAL DISCLAIMER: This NDA has been generated to comply with South African law as of 2024. However, legal requirements may change, and specific circumstances may require additional provisions. It is recommended to have this agreement reviewed by a qualified South African attorney before execution.
"""
    return contract

if __name__ == "__main__":
    main()